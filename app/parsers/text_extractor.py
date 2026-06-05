"""统一文档解析入口.

- PDF / DOCX / TXT 走同一个 TextExtractor 协议
- 对解析失败、加密、扫描件等异常做明确归类
- 仅负责"提取 + 清洗",不做 LLM 结构化
"""
from __future__ import annotations

import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional

from app.core.errors import ParseError, UnsupportedFileTypeError


class TextExtractor(ABC):
    """文本提取器抽象基类."""

    @abstractmethod
    def supports(self, suffix: str) -> bool: ...

    @abstractmethod
    def extract(self, file_path: str) -> str: ...


class DocxTextExtractor(TextExtractor):
    def supports(self, suffix: str) -> bool:
        return suffix.lower() == ".docx"

    def extract(self, file_path: str) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:  # pragma: no cover - 安装时已经包含
            raise ParseError("python-docx 未安装,无法解析 Word 文件") from exc

        try:
            doc = Document(file_path)
        except Exception as exc:  # noqa: BLE001
            raise ParseError(f"读取 DOCX 失败: {exc}") from exc

        parts: list[str] = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)


class PdfTextExtractor(TextExtractor):
    def supports(self, suffix: str) -> bool:
        return suffix.lower() == ".pdf"

    def extract(self, file_path: str) -> str:
        try:
            from pypdf import PdfReader  # 优先使用维护活跃的 pypdf
        except ImportError:  # 兼容 PyPDF2
            try:
                from PyPDF2 import PdfReader  # type: ignore
            except ImportError as exc:  # pragma: no cover
                raise ParseError("未安装 pypdf 或 PyPDF2,无法解析 PDF") from exc

        try:
            reader = PdfReader(file_path)
            if getattr(reader, "is_encrypted", False):
                raise ParseError("PDF 已加密,需要先解密")
            parts: list[str] = []
            for page in reader.pages:
                try:
                    text = page.extract_text() or ""
                except Exception as exc:  # noqa: BLE001
                    raise ParseError(f"PDF 页面解析失败: {exc}") from exc
                if text:
                    parts.append(text)
            return "\n".join(parts)
        except ParseError:
            raise
        except Exception as exc:  # noqa: BLE001
            raise ParseError(f"PDF 解析失败: {exc}") from exc


class TxtTextExtractor(TextExtractor):
    def supports(self, suffix: str) -> bool:
        return suffix.lower() in {".txt", ".md"}

    def extract(self, file_path: str) -> str:
        return Path(file_path).read_text(encoding="utf-8", errors="ignore")


_EXTRACTORS: list[TextExtractor] = [
    DocxTextExtractor(),
    PdfTextExtractor(),
    TxtTextExtractor(),
]


def extract_text(file_path: str) -> str:
    """根据文件后缀选择合适的提取器."""
    suffix = Path(file_path).suffix
    for extractor in _EXTRACTORS:
        if extractor.supports(suffix):
            raw = extractor.extract(file_path)
            return clean_text(raw)
    raise UnsupportedFileTypeError(f"不支持的文件类型: {suffix}")


_WHITESPACE_RE = re.compile(r"[\u00A0\u3000\t]+")
_MULTI_NEWLINE_RE = re.compile(r"\n{3,}")


def clean_text(text: str) -> str:
    """轻量清洗:统一空白、压缩连续换行."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WHITESPACE_RE.sub(" ", text)
    text = _MULTI_NEWLINE_RE.sub("\n\n", text)
    return text.strip()


def extract_text_from_bytes(data: bytes, suffix: str) -> str:
    """从内存字节直接解析,避免一定要先落盘."""
    import tempfile

    tmp: Optional[tempfile.NamedTemporaryFile] = None
    try:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        tmp.write(data)
        tmp.flush()
        tmp.close()
        return extract_text(tmp.name)
    finally:
        if tmp is not None:
            try:
                Path(tmp.name).unlink(missing_ok=True)
            except OSError:
                pass
